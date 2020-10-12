from spellchecker import SpellChecker
import re
from docx import Document
from docx.shared import RGBColor
import pandas as pd 

# Enter the file name and path
path=r'E:\GitHub\spell_corrections'
file_name='Assignment_Sampledata.txt'


spell = SpellChecker()
file1 = open(path+'\\'+file_name, 'r', encoding="utf8") 
Lines = file1.readlines()


class Text_Validation:
    def CheckError(Lines):
        count=1
        word_list=[]
        for line in Lines:
    
            res1 = re.sub(r'[^\w\s]', '', line)
            res = re.sub('[0-9]', '', res1)
    
        
            try:
                misspelled = spell.unknown(res.split())
                
                for word in misspelled:
                    word_list.append((count,word,spell.correction(word)))
                count=count+1
            
            except Exception as e:
                print(e)
         
        return word_list
    
    
    
file1.close()


class DocFile_generator: 
    file_name_doc='result.docx'        
    def CreateDocx(self,answer):
        document = Document()
        document.add_heading('Spell Correction', 0)
        
        table = document.add_table(rows=1, cols=3)
        heading_cells = table.rows[0].cells
        run = heading_cells[0].paragraphs[0].add_run('Line')
        run.font.bold = True
        run = heading_cells[1].paragraphs[0].add_run('word')
        run.font.bold = True
        run = heading_cells[2].paragraphs[0].add_run('correct_word')
        run.font.bold = True
        
        for Line, word, correct_word in answer:
            row_cells = table.add_row().cells
            my_cell = row_cells[0]
            row_cells[0].paragraphs[0].add_run(str(Line))
            row_cells[1].paragraphs[0].add_run(word)
            run=row_cells[2].paragraphs[0].add_run(correct_word)
            red = RGBColor(255, 0, 0)
            run.font.color.rgb = red
            run.font.bold = True

        document.save(path+'\\'+self.file_name_doc)
        
       
class ExcelFile_generator:
    file_name_Excel='result.xlsx' 
    def CreatExcel(self,answer):
            df = pd.DataFrame(columns=['Line', 'word', 'correct_word'])
            for count, word, correct_word in answer:
                df = df.append({'Line': count, 'word': word, 'correct_word': correct_word}, ignore_index=True)
                
            writer = pd.ExcelWriter(path+'\\'+self.file_name_Excel, engine='xlsxwriter')
            df.to_excel(writer, sheet_name="Sheet1", index=False)
            
            worksheet = writer.sheets['Sheet1']
            workbook = writer.book
            cell_format = workbook.add_format()
            cell_format.set_bold()
            cell_format.set_font_color('green')
            
            worksheet.set_column('C:C', None, cell_format)
            
            writer.close()
        
answer = Text_Validation.CheckError(Lines)

DocFile_generator().CreateDocx(answer)

ExcelFile_generator().CreatExcel(answer)
